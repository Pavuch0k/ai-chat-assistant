from fastapi import APIRouter, Depends, UploadFile, File, HTTPException
from sqlalchemy.orm import Session
from app.db.database import get_db
from app.models.db_models import Document
from app.services.knowledge_service import knowledge_service
from typing import List, Optional
from pydantic import BaseModel
from datetime import datetime
import os
import uuid

router = APIRouter(prefix="/api/admin/knowledge", tags=["knowledge"])

# Определяем путь для загрузок (локально или в Docker)
UPLOAD_DIR = os.getenv("UPLOAD_DIR", os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "uploads"))
os.makedirs(UPLOAD_DIR, exist_ok=True)

class DocumentResponse(BaseModel):
    id: int
    name: str
    file_type: Optional[str] = None
    file_size: Optional[int] = None
    created_at: datetime
    
    class Config:
        from_attributes = True

@router.get("", response_model=List[DocumentResponse])
async def get_documents(db: Session = Depends(get_db)):
    """Получить список всех документов"""
    documents = db.query(Document).order_by(Document.created_at.desc()).all()
    return documents

@router.post("/upload")
async def upload_documents(
    files: List[UploadFile] = File(...),
    db: Session = Depends(get_db)
):
    """Загрузить документы в базу знаний"""
    uploaded_docs = []
    
    for file in files:
        try:
            # Сохраняем файл
            file_ext = os.path.splitext(file.filename)[1]
            unique_filename = f"{uuid.uuid4()}{file_ext}"
            file_path = os.path.join(UPLOAD_DIR, unique_filename)
            
            content = await file.read()
            with open(file_path, "wb") as f:
                f.write(content)
            
            # Читаем текст из файла
            text = ""
            if file_ext.lower() == '.txt':
                try:
                    with open(file_path, "r", encoding="utf-8") as f:
                        text = f.read()
                except UnicodeDecodeError:
                    with open(file_path, "r", encoding="windows-1251") as f:
                        text = f.read()
            elif file_ext.lower() == '.pdf':
                try:
                    from pypdf import PdfReader
                    reader = PdfReader(file_path)
                    text = "\n".join([page.extract_text() for page in reader.pages])
                except Exception as e:
                    text = f"Не удалось извлечь текст из PDF: {file.filename} - {str(e)}"
            elif file_ext.lower() in ['.doc', '.docx']:
                try:
                    from docx import Document
                    # Открываем файл напрямую по пути
                    doc = Document(file_path)
                    paragraphs = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            paragraphs.append(paragraph.text)
                    # Также извлекаем текст из таблиц
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    paragraphs.append(cell.text)
                    text = "\n".join(paragraphs)
                except Exception as e:
                    import traceback
                    error_details = traceback.format_exc()
                    print(f"DOCX error details: {error_details}")
                    text = f"Не удалось извлечь текст из DOC/DOCX: {file.filename} - {str(e)}"
            else:
                text = f"Формат {file_ext} пока не поддерживается"
            
            # Создаем запись в БД
            document = Document(
                name=file.filename,
                file_path=file_path,
                file_type=file_ext,
                file_size=len(content)
            )
            db.add(document)
            db.commit()
            db.refresh(document)
            
            # Добавляем в векторную БД
            if text:
                knowledge_service.add_document(
                    text=text,
                    document_id=document.id,
                    metadata={"name": file.filename, "type": file_ext}
                )
            
            uploaded_docs.append(document)
        except Exception as e:
            print(f"Error uploading file {file.filename}: {e}")
            continue
    
    return {"message": f"Загружено документов: {len(uploaded_docs)}", "documents": uploaded_docs}

@router.delete("/{document_id}")
async def delete_document(document_id: int, db: Session = Depends(get_db)):
    """Удалить документ"""
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Документ не найден")
    
    # Удаляем из векторной БД
    knowledge_service.delete_document(document_id)
    
    # Удаляем файл
    if os.path.exists(document.file_path):
        os.remove(document.file_path)
    
    # Удаляем из БД
    db.delete(document)
    db.commit()
    
    return {"message": "Документ удален"}
